"""
DOCX Parser module.
Extracts raw text from DOCX files using python-docx.
"""

import docx
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a Word document (.docx) file.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract from tables in the docx
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        return ""
